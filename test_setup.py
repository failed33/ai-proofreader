"""
Test script to validate that all dependencies are installed correctly.
Run this before using the main proofreading script.
"""

def test_imports():
    """Test that all required packages can be imported."""
    print("Testing package imports...")
    
    try:
        import docx
        print("✅ python-docx: OK")
    except ImportError as e:
        print(f"❌ python-docx: FAILED - {e}")
        return False
    
    try:
        import tiktoken
        print("✅ tiktoken: OK")
    except ImportError as e:
        print(f"❌ tiktoken: FAILED - {e}")
        return False
    
    try:
        import openai
        print("✅ openai: OK")
    except ImportError as e:
        print(f"❌ openai: FAILED - {e}")
        return False
    
    try:
        import backoff
        print("✅ backoff: OK")
    except ImportError as e:
        print(f"❌ backoff: FAILED - {e}")
        return False
    
    try:
        import pydantic
        print("✅ pydantic: OK")
    except ImportError as e:
        print(f"❌ pydantic: FAILED - {e}")
        return False
    
    try:
        import tqdm
        print("✅ tqdm: OK")
    except ImportError as e:
        print(f"❌ tqdm: FAILED - {e}")
        return False
    
    try:
        from dotenv import load_dotenv
        print("✅ python-dotenv: OK")
    except ImportError as e:
        print(f"❌ python-dotenv: FAILED - {e}")
        return False
    
    return True

def test_model_availability():
    """Test that tiktoken can access the model we want to use."""
    print("\nTesting model availability...")
    
    try:
        import tiktoken
        models_to_test = ["gpt-4o-mini", "gpt-4", "gpt-3.5-turbo"]
        
        for model in models_to_test:
            try:
                enc = tiktoken.encoding_for_model(model)
                print(f"✅ {model}: Available")
            except KeyError:
                print(f"❌ {model}: Not available in tiktoken")
    except Exception as e:
        print(f"❌ Model test failed: {e}")
        return False
    
    return True

def test_env_setup():
    """Test environment variable setup."""
    print("\nTesting environment setup...")
    
    import os
    from dotenv import load_dotenv
    load_dotenv()
    
    api_key = os.getenv("OPENAI_API_KEY")
    if api_key:
        # Don't print the actual key for security
        print(f"✅ OPENAI_API_KEY: Set (length: {len(api_key)})")
        return True
    else:
        print("❌ OPENAI_API_KEY: Not set")
        print("   Please create a .env file with your OpenAI API key")
        return False

def test_docx_functionality():
    """Test basic Word document creation and reading."""
    print("\nTesting Word document functionality...")
    
    try:
        from docx import Document
        
        # Create a test document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        
        test_file = 'test_output.docx'
        doc.save(test_file)
        print("✅ Document creation: OK")
        
        # Read the test document
        doc2 = Document(test_file)
        paragraphs = [p.text.strip() for p in doc2.paragraphs if p.text.strip()]
        if len(paragraphs) > 0:
            print("✅ Document reading: OK")
        else:
            print("❌ Document reading: No paragraphs found")
            return False
        
        # Clean up
        import os
        os.remove(test_file)
        print("✅ Cleanup: OK")
        
        return True
        
    except Exception as e:
        print(f"❌ Word document test failed: {e}")
        return False

def main():
    """Run all tests."""
    print("🧪 Running setup validation tests...\n")
    
    tests = [
        ("Package Imports", test_imports),
        ("Model Availability", test_model_availability),
        ("Environment Setup", test_env_setup),
        ("Word Document Functionality", test_docx_functionality)
    ]
    
    all_passed = True
    
    for test_name, test_func in tests:
        try:
            result = test_func()
            if not result:
                all_passed = False
        except Exception as e:
            print(f"❌ {test_name}: Exception - {e}")
            all_passed = False
        
        print()  # Add spacing between tests
    
    print("=" * 50)
    if all_passed:
        print("🎉 All tests passed! You're ready to use the proofreading script.")
    else:
        print("⚠️  Some tests failed. Please fix the issues before proceeding.")
        print("\nNext steps:")
        print("1. Install missing packages: pip install -r requirements.txt")
        print("2. Set up your OpenAI API key in a .env file")
        print("3. Re-run this test script")

if __name__ == "__main__":
    main() 