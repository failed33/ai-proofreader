"""
Helper script to create a sample Word document with deliberate errors for testing the proofreader.
"""

from docx import Document

def create_sample_document():
    """Create a sample Word document with intentional errors for testing."""
    doc = Document()
    doc.add_heading('Sample Document for AI Proofreading', 0)
    
    # Sample paragraphs with various types of errors
    sample_paragraphs = [
        "This is the frist paragraph of the sample document. It contians some speling errors and grammer mistakes that the AI proofreader should be able to catch and fix.",
        
        "The secound paragraph has different types of issues. For example, it have subject-verb disagreement, missing commas and some akward phrasing that could be improved.",
        
        "In this third paragrpah, we're testing whether the AI can identify redundant words words and improve the overall clarity and conciseness of the writting.",
        
        "The forth paragraph test the AI's ability to catch more subtle errors like incorect word usage (e.g., 'affect' vs 'effect'), improper capitalization, and run-on sentences that should really be broken up into multiple sentences for better readability and comprehension.",
        
        "Finally the last paragraph deliberately omits punctuation at the end It also has some very long sentences that might benefit from being shortened and simplified for better clarity and reader comprehension which is always important in good writing"
    ]
    
    for paragraph in sample_paragraphs:
        doc.add_paragraph(paragraph)
    
    doc.save('my_draft.docx')
    print("✅ Sample document 'my_draft.docx' created successfully!")
    print("This document contains intentional errors for testing the AI proofreader.")

if __name__ == "__main__":
    create_sample_document() 