"""
AI Proofreader for Word Documents (Modular Version)

This script takes a Word document (.docx), sends each paragraph to OpenAI for proofreading,
and generates a comprehensive report with original text, corrections, and feedback.
"""

import json
from typing import List
import tiktoken
import openai
import backoff
from pydantic import BaseModel, ValidationError
from docx import Document
from tqdm import tqdm

# Import configuration
from config import (
    OPENAI_API_KEY, OPENAI_MODEL, TEMPERATURE,
    MAX_TOKENS_PER_API_REQUEST, INPUT_DOCX_PATH, OUTPUT_DOCX_PATH,
    MAX_RETRY_TIME_SECONDS, SYSTEM_PROMPT, USER_PROMPT_TEMPLATE
)

# Initialize OpenAI client
client = openai.OpenAI(api_key=OPENAI_API_KEY)


# --- Pydantic Models ---
class Edit(BaseModel):
    original: str
    corrected: str
    feedback: List[str]


# --- Core Functions ---
def load_paragraphs(docx_path: str) -> List[str]:
    """Load paragraphs from a Word document."""
    import os
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Input file '{docx_path}' not found.")
    
    doc = Document(docx_path)
    raw_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    print(f"Loaded {len(raw_paragraphs)} paragraphs from '{docx_path}'.")
    return raw_paragraphs


def prepare_segments_for_api(paragraphs: List[str], model_name: str, max_tokens_for_request: int) -> List[dict]:
    """Prepare paragraph segments for API processing, checking token limits."""
    enc = tiktoken.encoding_for_model(model_name)
    segments_to_process = []
    
    # Estimate tokens used by the prompt template (excluding the paragraph text itself)
    prompt_overhead = len(enc.encode(USER_PROMPT_TEMPLATE.replace("{paragraph_text}", ""))) + \
                      len(enc.encode(SYSTEM_PROMPT)) + 200  # 200 as a buffer for JSON structure, roles etc.

    for i, para_text in enumerate(paragraphs):
        num_tokens = len(enc.encode(para_text))
        if num_tokens + prompt_overhead > max_tokens_for_request:
            print(f"Warning: Paragraph {i+1} is too long ({num_tokens} tokens) and might exceed API limits.")
            print(f"  Preview: {para_text[:100]}...")
            segments_to_process.append({"id": i, "text": para_text, "is_long": True})
        else:
            segments_to_process.append({"id": i, "text": para_text, "is_long": False})
    
    print(f"Prepared {len(segments_to_process)} segments for API processing.")
    return segments_to_process


@backoff.on_exception(
    backoff.expo,
    (openai.RateLimitError, openai.APIConnectionError, openai.APITimeoutError),
    max_time=MAX_RETRY_TIME_SECONDS
)
def get_proofread_data(paragraph_text: str) -> dict:
    """Call OpenAI API to proofread a paragraph and return structured response."""
    user_content = USER_PROMPT_TEMPLATE.format(paragraph_text=paragraph_text)
    
    try:
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_content}
            ],
            temperature=TEMPERATURE,
            response_format={"type": "json_object"}
        )
        return json.loads(response.choices[0].message.content)
    except json.JSONDecodeError as e:
        print(f"JSON decode error for paragraph: '{paragraph_text[:50]}...': {e}")
        return {"original": paragraph_text, "corrected": paragraph_text, "feedback": [f"JSON Error: {str(e)}"]}
    except Exception as e:
        print(f"Error during API call for paragraph: '{paragraph_text[:50]}...': {e}")
        return {"original": paragraph_text, "corrected": paragraph_text, "feedback": [f"API Error: {str(e)}"]}


def process_segments(segments_to_process: List[dict]) -> List[Edit]:
    """Process all segments through the API and validate responses."""
    results = []
    
    for segment_data in tqdm(segments_to_process, desc="Proofreading paragraphs"):
        paragraph_text = segment_data["text"]
        
        try:
            raw_json_response = get_proofread_data(paragraph_text)
            validated_data = Edit.model_validate(raw_json_response)
            results.append(validated_data)
        except ValidationError as e:
            print(f"Pydantic Validation Error for paragraph: '{paragraph_text[:50]}...'\n{e}")
            results.append(Edit(
                original=paragraph_text, 
                corrected=paragraph_text, 
                feedback=[f"Validation Error: {str(e)}"]
            ))
        except Exception as e:
            print(f"General Error processing paragraph: '{paragraph_text[:50]}...': {e}")
            results.append(Edit(
                original=paragraph_text, 
                corrected=paragraph_text, 
                feedback=[f"Processing Error: {str(e)}"]
            ))
    
    return results


def create_output_document(proofread_results: List[Edit], output_path: str) -> None:
    """Create a Word document with proofreading results."""
    output_doc = Document()
    output_doc.add_heading("AI Proofreading Report", level=0)
    
    # Add summary information
    total_paragraphs = len(proofread_results)
    corrected_paragraphs = sum(1 for edit in proofread_results if edit.original != edit.corrected)
    
    summary = output_doc.add_paragraph()
    summary.add_run("Summary: ").bold = True
    summary.add_run(f"Processed {total_paragraphs} paragraphs. {corrected_paragraphs} paragraphs had corrections.")
    
    output_doc.add_paragraph("=" * 50)
    
    for i, edit_item in enumerate(proofread_results):
        # Original paragraph
        output_doc.add_heading(f"Paragraph {i+1}: Original", level=2)
        original_para = output_doc.add_paragraph(edit_item.original)
        original_para.style = 'Quote'
        
        # Corrected paragraph
        output_doc.add_heading(f"Paragraph {i+1}: Corrected", level=2)
        corrected_para = output_doc.add_paragraph(edit_item.corrected)
        
        # Highlight if there are changes
        if edit_item.original != edit_item.corrected:
            corrected_para.runs[0].font.highlight_color = 3  # Yellow highlight
        
        # Feedback
        output_doc.add_heading(f"Paragraph {i+1}: Feedback", level=2)
        if edit_item.feedback:
            feedback_items = [item.strip() for item in edit_item.feedback if item.strip()]
            if feedback_items:
                for item in feedback_items:
                    # Remove leading bullet if AI already added one
                    clean_item = item[2:].strip() if item.startswith("- ") else item
                    if clean_item:
                        output_doc.add_paragraph(clean_item, style='List Bullet')
            else:
                output_doc.add_paragraph("No specific feedback provided.")
        else:
            output_doc.add_paragraph("No feedback provided.")
        
        # Separator
        output_doc.add_paragraph("-" * 30)
    
    try:
        output_doc.save(output_path)
        print(f"✅ Output document saved to '{output_path}'")
    except Exception as e:
        print(f"❌ Error saving output document: {e}")


def main():
    """Main function to orchestrate the proofreading process."""
    print("🚀 Starting AI proofreading script...")
    print(f"📖 Input file: {INPUT_DOCX_PATH}")
    print(f"📝 Output file: {OUTPUT_DOCX_PATH}")
    print(f"🤖 Using model: {OPENAI_MODEL}")
    print("-" * 50)
    
    try:
        # Load Word Document
        raw_paragraphs = load_paragraphs(INPUT_DOCX_PATH)
        if not raw_paragraphs:
            print("❌ No paragraphs found in the document. Exiting.")
            return
        
        # Token Management & Segment Preparation
        segments_to_process = prepare_segments_for_api(
            raw_paragraphs,
            OPENAI_MODEL,
            MAX_TOKENS_PER_API_REQUEST
        )
        if not segments_to_process:
            print("❌ No segments to process after token management. Exiting.")
            return
        
        # Process segments through API
        proofread_results = process_segments(segments_to_process)
        
        # Generate Output Word Document
        create_output_document(proofread_results, OUTPUT_DOCX_PATH)
        
        print("✅ Proofreading script finished successfully!")
        
        # Print summary statistics
        total_paragraphs = len(proofread_results)
        corrected_paragraphs = sum(1 for edit in proofread_results if edit.original != edit.corrected)
        print(f"📊 Summary: {corrected_paragraphs}/{total_paragraphs} paragraphs were corrected.")
        
    except Exception as e:
        print(f"❌ Script failed with error: {e}")
        raise


if __name__ == "__main__":
    main() 